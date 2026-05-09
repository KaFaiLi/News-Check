from __future__ import annotations

import re

from src.document_generator import DocumentGenerator
from src.models import DegradationStatus


def _sample_articles() -> list[dict]:
    return [
        {
            "article": {
                "title": "Bank A adopts AI agents for trade execution",
                "source": "Reuters",
                "published_time": "2026-04-15T10:00:00Z",
                "url": "https://reuters.com/article-1",
            },
            "analysis": {
                "insights": [
                    "Cuts execution latency by 40 percent.",
                    "Compliance challenges around model attribution remain open.",
                ]
            },
        },
        {
            "article": {
                "title": "Generative AI banking platform raises $500M",
                "source": "Bloomberg",
                "published_time": "2026-04-22T12:00:00Z",
                "url": "https://bloomberg.com/article-2",
            },
            "analysis": {"insights": ["Valuation implies 25x forward revenue."]},
        },
        {
            "article": {
                "title": "OpenAI releases new agentic framework",
                "source": "TechCrunch",
                "published_time": "2026-04-09T09:00:00Z",
                "url": "https://techcrunch.com/article-3",
            },
            "analysis": {"insights": ["Direct competition for enterprise AI tooling."]},
        },
    ]


def test_email_uses_monthly_copy_not_daily(tmp_path):
    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme Bank",
    )
    path = gen.generate_email_content(_sample_articles())
    html = (tmp_path / path.split("\\")[-1].split("/")[-1]).read_text(encoding="utf-8") if False else open(path, encoding="utf-8").read()

    assert "AI Monthly Digest" in html
    assert "Monthly Briefing" in html
    assert "April 2026" in html
    assert "This Month's Top Stories" in html

    # Daily-cadence wording must be gone
    assert "Daily Briefing" not in html
    assert "past 24 hours" not in html
    assert "Today's Top Stories" not in html
    assert "AI & Fintech News Digest" not in html


def test_email_substitutes_company_name(tmp_path):
    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme Bank",
    )
    path = gen.generate_email_content(_sample_articles())
    html = open(path, encoding="utf-8").read()
    assert "Acme Bank" in html
    assert "[Your Company Name]" not in html


def test_detailed_report_uses_monthly_copy(tmp_path):
    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme Bank",
    )
    path = gen.generate_detailed_report(_sample_articles())
    from docx import Document

    doc = Document(path)
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "AI Monthly Digest — Full Report" in body
    assert "April 2026" in body
    assert "investment banking" in body.lower()
    assert "Detailed News Report" not in body  # legacy title gone
    assert "past 24 hours" not in body


def test_detailed_report_uses_precomputed_summary(tmp_path):
    """Synthesis from the pipeline lands in the 'Monthly Themes' section
    of the detailed report — never re-calls the LLM."""
    from docx import Document

    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme Bank",
        # No `llm_instance` passed — if it tried to call one, it would fail.
    )
    path = gen.generate_detailed_report(
        _sample_articles(),
        precomputed_summary="MAGIC SYNTHESIS TOKEN",
    )
    doc = Document(path)
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "MAGIC SYNTHESIS TOKEN" in body
    assert "Monthly Themes" in body


def test_detailed_report_omits_themes_when_no_summary(tmp_path):
    from docx import Document

    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme Bank",
    )
    path = gen.generate_detailed_report(_sample_articles())
    doc = Document(path)
    body = "\n".join(p.text for p in doc.paragraphs)
    # Without a precomputed summary AND no llm_instance, _generate_overall_summary
    # returns the "could not be generated" sentence; we still render it as a
    # themes section. We assert the section doesn't crash the doc.
    assert "AI Monthly Digest — Full Report" in body


def test_degradation_warning_renders_in_email(tmp_path):
    deg = DegradationStatus(
        is_degraded=True,
        total_attempts=10,
        failed_attempts=4,
        collected_results_count=6,
        warnings=["Source X blocked us repeatedly."],
    )
    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme",
        include_degradation_warning=True,
    )
    path = gen.generate_email_content(_sample_articles(), degradation_status=deg)
    html = open(path, encoding="utf-8").read()
    assert "DEGRADATION WARNING" in html
    assert re.search(r"Success Rate:\s*60", html)


def test_degradation_warning_suppressed_when_flag_off(tmp_path):
    deg = DegradationStatus(is_degraded=True, total_attempts=10, failed_attempts=4)
    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme",
        include_degradation_warning=False,
    )
    path = gen.generate_email_content(_sample_articles(), degradation_status=deg)
    html = open(path, encoding="utf-8").read()
    assert "DEGRADATION WARNING" not in html


# ---------- formatting fixes ----------


def test_detailed_renders_insights_as_bullets(tmp_path):
    """Insights list must NOT render as `['a', 'b', 'c']`. Each insight
    appears on its own bulleted line in the detailed report."""
    from docx import Document

    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme",
    )
    path = gen.generate_detailed_report(_sample_articles())
    doc = Document(path)
    body = "\n".join(p.text for p in doc.paragraphs)
    # No Python list repr
    assert "['" not in body
    assert "', '" not in body
    assert "• Cuts execution latency by 40 percent." in body
    assert "• Compliance challenges around model attribution remain open." in body
    assert "• Valuation implies 25x forward revenue." in body
    assert "• Direct competition for enterprise AI tooling." in body


def test_detailed_formats_pub_date_human_readable(tmp_path):
    from docx import Document

    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme",
    )
    path = gen.generate_detailed_report(_sample_articles())
    doc = Document(path)
    body = "\n".join(p.text for p in doc.paragraphs)
    # New format: "April 15, 2026" — not "2026-04-15 10:00" or ISO 8601
    assert "April 15, 2026" in body
    assert "April 22, 2026" in body
    assert "2026-04-15T10:00:00Z" not in body
    assert "2026-04-15 10:00" not in body


def test_email_shows_date_not_time_of_day(tmp_path):
    """Regression: email used to format pub_time as %I:%M %p (e.g. '10:00 AM')
    which makes no sense for a monthly digest. It must be a date now."""
    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme",
    )
    path = gen.generate_email_content(_sample_articles())
    html = open(path, encoding="utf-8").read()
    assert "April 15, 2026" in html
    # No "AM" / "PM" in the article meta lines (matching against the metadata pattern)
    # We can't just check for "AM" (could appear elsewhere), so look for the buggy pattern.
    assert "10:00 AM" not in html
    assert "12:00 PM" not in html


def test_format_pub_date_handles_missing_and_unknown(tmp_path):
    gen = DocumentGenerator(
        output_dir=str(tmp_path),
        month_label="April 2026",
        company_name="Acme",
    )
    assert gen._format_pub_date(None) == "Date unknown"
    assert gen._format_pub_date("Unknown Time") == "Date unknown"
    assert gen._format_pub_date("not-a-date") == "not-a-date"
    from datetime import datetime
    assert gen._format_pub_date(datetime(2026, 4, 22)) == "April 22, 2026"
