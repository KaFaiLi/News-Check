"""Word + HTML email generation for the monthly AI digest.

Layout, styling, fonts, and the company-red `#E9041E` accent are preserved
from the legacy template. Copy strings are tuned for an investment-banking
audience reading a monthly briefing.
"""

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import AzureChatOpenAI


class DocumentGenerator:
    def __init__(
        self,
        output_dir: Path | str = "Output",
        llm_instance: AzureChatOpenAI | None = None,
        *,
        month_label: str = "",
        company_name: str = "[Your Company Name]",
        include_degradation_warning: bool = True,
    ):
        self.output_dir = Path(output_dir)
        self.llm = llm_instance
        self.month_label = month_label or datetime.now().strftime("%B %Y")
        self.company_name = company_name
        self.include_degradation_warning = include_degradation_warning
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.summary_prompt = ChatPromptTemplate.from_messages(
            [
                (
                    "system",
                    "You are a senior research analyst at a global investment bank. "
                    "Synthesize the key AI / AI-banking / AI-agents themes from the "
                    "supplied articles into a brief, IB-relevant digest opening.",
                ),
                (
                    "user",
                    "Provide a concise 3-4 sentence overall synthesis for {month_label} "
                    "based on the following articles:\n\n{article_summaries}\n\nOverall Summary:",
                ),
            ]
        )
        if self.llm:
            self.summary_chain = self.summary_prompt | self.llm
        else:
            self.summary_chain = None

    def _set_doc_margins(self, document):
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_styled_paragraph(
        self, document, text, size=11, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        paragraph.paragraph_format.space_after = Pt(6)
        return paragraph

    @staticmethod
    def _format_pub_date(pub_time_str: Any) -> str:
        """Render a `published_time` field as 'April 22, 2026'.

        Accepts ISO-8601 strings (with or without 'Z'), datetime objects,
        or 'Unknown Time' / falsy values. Falls back to the raw input
        if parsing fails."""
        if not pub_time_str or pub_time_str == "Unknown Time":
            return "Date unknown"
        if isinstance(pub_time_str, datetime):
            return pub_time_str.strftime("%B %d, %Y")
        try:
            dt = datetime.fromisoformat(str(pub_time_str).replace("Z", "+00:00"))
            return dt.strftime("%B %d, %Y")
        except (ValueError, TypeError):
            return str(pub_time_str)

    @staticmethod
    def _normalize_insights(insights: Any) -> list[str]:
        """Coerce any of {list[str], '•'-separated string, plain string, None}
        to a clean list of bullet strings. Empty input → empty list."""
        if isinstance(insights, list):
            return [str(b).strip() for b in insights if str(b).strip()]
        if isinstance(insights, str) and insights.strip():
            text = insights.strip()
            if "•" in text:
                return [p.strip() for p in text.split("•") if p.strip()]
            return [text]
        return []

    def _render_insight_bullets(self, document, insights: Any) -> None:
        """Render insights as bulleted paragraphs in the docx."""
        bullets = self._normalize_insights(insights)
        if not bullets:
            self._add_styled_paragraph(
                document, "No analysis available for this article.", size=11
            )
            return
        for b in bullets:
            p = document.add_paragraph()
            run = p.add_run(f"• {b}")
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(3)

    @staticmethod
    def _render_insights_html(insights: Any) -> str:
        """HTML version of bullet rendering for the email path."""
        bullets = DocumentGenerator._normalize_insights(insights)
        if not bullets:
            return (
                "<p style='margin: 10px 0 0 0;'>"
                "No analysis available for this article.</p>"
            )
        items = "".join(
            f"<li style='margin-bottom: 6px;'>{b}</li>" for b in bullets
        )
        return (
            f"<ul style='margin: 10px 0 0 0; padding-left: 20px;'>{items}</ul>"
        )

    def _format_degradation_summary(
        self, degradation_status: Any, fallback_count: int
    ) -> str:
        """Single-line degradation summary used by docx, html, and xlsx."""
        rate = getattr(degradation_status, "success_rate", 0.0)
        failed = getattr(degradation_status, "failed_attempts", 0)
        total = getattr(degradation_status, "total_attempts", 0)
        collected = getattr(degradation_status, "collected_results_count", 0) or fallback_count
        return (
            f"Success Rate: {rate:.1%} | "
            f"Failed Attempts: {failed}/{total} | "
            f"Collected Results: {collected} articles"
        )

    def _generate_overall_summary(
        self,
        top_articles: list[dict],
        precomputed_summary: str | None = None,
    ) -> str:
        """Return the digest's opening synthesis.

        If `precomputed_summary` is provided (the standard path from the
        pipeline), it is returned directly. Falls back to the local LLM
        chain only when the caller does not have one and an LLM was
        configured at construction time.
        """
        if precomputed_summary and precomputed_summary.strip():
            return precomputed_summary.strip()

        if not self.llm or not self.summary_chain:
            return "Overall summary could not be generated (no precomputed summary and no LLM configured)."

        summary_input = []
        for item in top_articles:
            title = item.get("article", {}).get("title", "No Title")
            insight = item.get("analysis", {}).get("insights", "N/A")
            insight_str = str(insight) if insight else "N/A"
            summary_input.append(f"Title: {title}\nInsight: {insight_str}\n---")

        if not summary_input:
            return "No articles available to generate a summary."

        try:
            print("Generating overall summary with LLM (fallback path)...")
            response = self.summary_chain.invoke(
                {
                    "article_summaries": "\n".join(summary_input),
                    "month_label": self.month_label,
                }
            )
            return str(response.content)
        except Exception as e:
            print(f"Error generating overall summary with LLM: {e}")
            return f"Overall summary could not be generated due to an error: {e}"

    def generate_detailed_report(
        self,
        top_articles: list[dict],
        degradation_status: Any | None = None,
        precomputed_summary: str | None = None,
    ):
        """Generates a detailed Word document report including a table of contents."""
        document = Document()
        self._set_doc_margins(document)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"detailed_news_report_{timestamp}.docx"
        filepath = self.output_dir / filename

        self._add_styled_paragraph(
            document,
            "AI Monthly Digest — Full Report",
            size=16,
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        self._add_styled_paragraph(
            document,
            f"{self.month_label} • Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            size=10,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )

        if (
            self.include_degradation_warning
            and degradation_status
            and getattr(degradation_status, "is_degraded", False)
        ):
            document.add_paragraph()
            self._add_styled_paragraph(
                document, "[WARN] DEGRADATION WARNING", size=12, bold=True
            )
            warning_text = (
                "This report was generated under degraded conditions due to sustained blocking/errors.\n"
                + self._format_degradation_summary(degradation_status, len(top_articles))
            )
            self._add_styled_paragraph(document, warning_text, size=10)
            if degradation_status.warnings:
                for warning in degradation_status.warnings[:5]:
                    self._add_styled_paragraph(document, f"  • {warning}", size=9)

        document.add_paragraph()

        # Monthly themes synthesis (uses precomputed summary from the pipeline,
        # falls back to a local LLM call if provided one at construction time).
        overall_summary_text = self._generate_overall_summary(
            top_articles, precomputed_summary=precomputed_summary
        )
        if overall_summary_text:
            self._add_styled_paragraph(document, "Monthly Themes", size=14, bold=True)
            self._add_styled_paragraph(document, overall_summary_text, size=11)
            document.add_paragraph()

        # Add Table of Contents
        self._add_styled_paragraph(document, "Table of Contents", size=14, bold=True)
        for i, item in enumerate(top_articles, 1):
            title = item.get("article", {}).get("title", "No Title Provided")
            toc_paragraph = self._add_styled_paragraph(
                document, f"{i}. {title}", size=11
            )
            toc_paragraph.paragraph_format.left_indent = Inches(0.25)
        document.add_paragraph()  # Add spacing after TOC

        self._add_styled_paragraph(
            document,
            f"The following are the top {len(top_articles)} AI stories for "
            f"{self.month_label}, curated for investment banking review.",
            size=11,
        )
        self._add_styled_paragraph(
            document,
            "For deeper analysis of each story, refer to the per-article insights below or the source links.",
            size=11,
        )
        document.add_paragraph()

        self._add_styled_paragraph(document, "Top Stories — Detail", size=14, bold=True)

        for i, item in enumerate(top_articles, 1):
            article = item.get("article", {})
            analysis = item.get("analysis", {})
            title = article.get("title", "No Title Provided")
            source = article.get("source", "Unknown Source")
            pub_time = self._format_pub_date(article.get("published_time"))
            url = article.get("url", "#")

            self._add_styled_paragraph(document, f"{i}. {title}", size=12, bold=True)
            self._add_styled_paragraph(
                document, f"{source}  •  {pub_time}", size=10
            )
            self._render_insight_bullets(document, analysis.get("insights"))
            p = self._add_styled_paragraph(document, "Link: ", size=10)
            p.add_run(url).font.size = Pt(10)
            document.add_paragraph()

        document.save(filepath)
        return str(filepath)

    def generate_email_content(
        self,
        top_articles: list[dict],
        degradation_status: Any | None = None,
    ) -> str:
        """Generates HTML content ready for copying into Outlook email."""

        month_label = self.month_label

        html_content = f"""\
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI Monthly Digest — {month_label}</title>
    <style>
        body {{
            margin: 0;
            padding: 0;
            background-color: #f0f2f5; /* Light grey background for the email client viewport */
            font-family: Arial, sans-serif;
            -webkit-text-size-adjust: 100%; /* Prevent iOS font scaling */
            -ms-text-size-adjust: 100%; /* Prevent Windows Mobile font scaling */
        }}
        table {{ /* Outlook fix for unwanted spacing */
            border-collapse: collapse;
            mso-table-lspace: 0pt;
            mso-table-rspace: 0pt;
        }}
        .email-container {{
            max-width: 700px;
            margin: 20px auto; /* Centering and top/bottom margin */
            background-color: #ffffff; /* White background for the content area */
            border: 1px solid #dddddd;
            font-family: Arial, sans-serif;
            font-size: 11pt;
            color: #333333;
        }}
        .header {{
            padding: 25px 30px;
            border-bottom: 4px solid #E9041E; /* Company red accent */
            background-color: #ffffff;
        }}
        .header h1 {{
            font-size: 24pt;
            color: #1A1A1A; /* Darker, professional black */
            margin: 0 0 5px 0;
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
            font-weight: bold;
        }}
        .header p {{
            font-size: 11pt;
            color: #555555;
            margin: 0;
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        }}
        .intro-section {{
            padding: 25px 30px;
            background-color: #f8f8f8; /* Very light grey for intro box */
        }}
        .intro-section p {{
            margin: 0;
            color: #333333;
            line-height: 1.6;
            font-size: 11pt;
        }}
        .stories-title-section {{
            padding: 25px 30px 15px 30px; /* Added more top padding */
        }}
        .stories-title-section h3 {{
            font-size: 17pt;
            color: #1A1A1A;
            margin: 0;
            padding-bottom: 8px;
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
            font-weight: bold;
        }}
        .article-item {{
            padding: 0px 30px 25px 30px; /* Padding around each article block */
        }}
        .article-table {{
            width: 100%;
            border-collapse: collapse;
            border: 1px solid #e0e0e0; /* Softer border for article box */
        }}
        .article-accent-cell {{
            width: 6px; /* Width of the red accent bar */
            background-color: #E9041E;
            font-size: 1px; /* Fix for some clients adding height */
            line-height: 1px; /* Fix for some clients adding height */
        }}
        .article-content-cell {{
            padding: 20px;
        }}
        .article-title {{ /* Encapsulating paragraph for title */
            margin: 0 0 8px 0;
        }}
        .article-title a {{
            font-size: 14pt;
            color: #222222; /* Dark, near-black for article titles */
            font-weight: bold;
            text-decoration: none;
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        }}
        .article-title a:hover {{
            text-decoration: underline;
        }}
        .article-meta {{
            margin: 0 0 12px 0; /* Adjusted margin */
            font-size: 10pt;
            color: #555555;
            font-family: Arial, sans-serif;
        }}
        .article-insights {{
            font-size: 11pt;
            color: #333333;
            line-height: 1.6;
            font-family: Arial, sans-serif;
        }}
        .article-insights ul {{
            margin: 10px 0 0 0; /* Top margin for ul */
            padding-left: 20px; /* Indent for bullet points */
        }}
        .article-insights li {{
            margin-bottom: 6px; /* Spacing between list items */
        }}
        .article-insights p {{ /* Paragraphs within insights */
            margin: 10px 0 0 0;
        }}
        .read-more-link a {{
            font-size: 10.5pt;
            color: #444444; /* Dark grey for the link */
            text-decoration: none;
            font-weight: bold;
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        }}
        .read-more-link a:hover {{
            text-decoration: underline;
        }}
        .footer {{
            padding: 25px 30px;
            border-top: 1px solid #dddddd;
            background-color: #f0f0f0; /* Light grey for footer */
            text-align: center;
        }}
        .footer p {{
            margin: 0 0 8px 0; /* Increased bottom margin */
            font-size: 9.5pt; /* Slightly adjusted for readability */
            color: #666666;
            line-height: 1.5;
        }}
        .footer p:last-child {{
            margin-bottom: 0;
        }}
        /* Alternating background colors for articles */
        .bg-white {{ background-color: #ffffff; }}
        .bg-lightgrey {{ background-color: #f9f9f9; }}

        /* Responsive considerations */
        @media screen and (max-width: 600px) {{
            .email-container {{
                width: 100% !important;
                margin: 0 auto !important;
                border-left: none !important;
                border-right: none !important;
            }}
            .header, .intro-section, .stories-title-section, .article-item, .footer {{
                padding-left: 20px !important;
                padding-right: 20px !important;
            }}
            .header h1 {{ font-size: 20pt !important; }}
            .stories-title-section h3 {{ font-size: 15pt !important; }}
            .article-title a {{ font-size: 13pt !important; }}
        }}
    </style>
</head>
<body>
    <table class="email-container" width="700" align="center" cellpadding="0" cellspacing="0" role="presentation" style="width: 700px;">
        <!-- Header -->
        <tr>
            <td class="header">
                <h1>AI Monthly Digest</h1>
                <p>Monthly Briefing &bull; {month_label}</p>
                    </td>
                </tr>

        <!-- Intro Section -->
        <tr>
            <td class="intro-section">
                <p>
                    Welcome to your monthly AI briefing, curated for investment banking professionals.
                    This digest highlights the most material developments in AI, AI in banking and finance,
                    and AI agents from {month_label}. Click any article title to read the full story.
                </p>"""

        if (
            self.include_degradation_warning
            and degradation_status
            and getattr(degradation_status, "is_degraded", False)
        ):
            summary_line = self._format_degradation_summary(
                degradation_status, len(top_articles)
            )
            html_content += f"""
                <div style="margin-top: 15px; padding: 15px; background-color: #fff3cd; border-left: 4px solid #ffc107;">
                    <p style="margin: 0; color: #856404; font-weight: bold; font-size: 11pt;">
                        [WARN] DEGRADATION WARNING
                    </p>
                    <p style="margin: 5px 0 0 0; color: #856404; font-size: 10pt;">
                        This report was generated under degraded conditions due to sustained blocking/errors.<br>
                        {summary_line}
                    </p>
                </div>"""

        html_content += """
                        </p>
                    </td>
                </tr>

        <!-- Spacer row for visual separation -->
        <tr><td style="height:10px; background-color: #ffffff; font-size: 1px; line-height: 1px;">&nbsp;</td></tr>

        <!-- Top Stories Title -->
        <tr>
            <td class="stories-title-section">
                <h3>This Month's Top Stories</h3>
                    </td>
                </tr>
        """

        # Add top 3 articles with new styling
        for i, item in enumerate(top_articles[:3], 1):
            article = item.get("article", {})
            analysis = item.get("analysis", {})

            title = article.get("title", "No Title Provided")
            source = article.get("source", "Unknown Source")
            pub_time = self._format_pub_date(article.get("published_time"))
            url = article.get("url", "#")

            processed_insights_html = self._render_insights_html(analysis.get("insights"))
            content_bg_class = "bg-lightgrey" if i % 2 == 0 else "bg-white"

            html_content += f"""\
        <!-- Article {i} -->
        <tr>
            <td class="article-item">
                <table class="article-table" cellpadding="0" cellspacing="0" role="presentation">
                    <tr>
                        <td class="article-accent-cell" style="width: 6px; background-color: #E9041E; font-size: 1px; line-height: 1px;">&nbsp;</td>
                        <td class="article-content-cell {content_bg_class}">
                            <p class="article-title">
                                <a href="{url}" target="_blank">{i}. {title}</a>
                            </p>
                            <p class="article-meta">
                                {source} &bull; {pub_time}
                            </p>
                            <div class="article-insights">
                                {processed_insights_html}
                            </div>
                            <p class="read-more-link" style="margin-top: 12px; margin-bottom: 0;">
                                <a href="{url}" target="_blank">Read Full Article &rarr;</a>
                            </p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
            """

        html_content += f"""\
        <!-- Spacer row before footer -->
        <tr><td style="height:10px; background-color: #ffffff; font-size: 1px; line-height: 1px;">&nbsp;</td></tr>

        <!-- Footer -->
        <tr>
            <td class="footer">
                <p>
                    This AI Monthly Digest is automatically generated for investment banking colleagues.
                    The full {len(top_articles)}-story report is attached as a Word document.
                </p>
                <p>
                    For feedback or inquiries, please reply to this email.
                </p>
                <p style="margin-top:15px; color: #888888;">&copy; {datetime.now().year} {self.company_name}. All rights reserved.</p>
                                </td>
                            </tr>
    </table> <!-- End of email-container table -->
</body>
</html>
"""

        timestamp_file = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"email_content_{timestamp_file}.html"
        filepath = self.output_dir / filename

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html_content)

        return str(filepath)

    def generate_articles_xlsx(
        self,
        scored_articles: list[Any],
        degradation_status: Any | None = None,
    ) -> str:
        """Write the selected scored articles to an Excel workbook.

        `scored_articles` is the list of `ScoredArticle` objects from the
        pipeline (so we can dump scores), but the method is duck-typed and
        will accept anything exposing `.article` and `.analysis`.
        """
        from openpyxl import Workbook

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"news_articles_{timestamp}.xlsx"
        filepath = self.output_dir / filename

        wb = Workbook()
        ws = wb.active
        ws.title = "Top Articles"
        ws.append(
            [
                "Rank",
                "Title",
                "Source",
                "Domain",
                "Published",
                "URL",
                "Final Score",
                "Topic Relevance",
                "Cross-Source Consensus",
                "Source Tier",
                "Source Tier Multiplier",
                "Recency",
                "Insights",
            ]
        )
        for i, sa in enumerate(scored_articles, 1):
            article = getattr(sa, "article", None)
            analysis = getattr(sa, "analysis", None)
            if article is None:
                continue
            insights = []
            if analysis and getattr(analysis, "insights", None):
                insights = analysis.insights
            pub_time = getattr(article, "published_time", None)
            ws.append(
                [
                    i,
                    getattr(article, "title", ""),
                    getattr(article, "source", ""),
                    getattr(article, "domain", ""),
                    pub_time.isoformat() if pub_time else "",
                    str(getattr(article, "url", "")),
                    round(float(getattr(sa, "final_score", 0.0)), 4),
                    round(float(getattr(sa, "topic_relevance", 0.0)), 4),
                    round(float(getattr(sa, "cross_source_consensus", 0.0)), 4),
                    int(getattr(sa, "source_tier", 3)) if getattr(sa, "source_tier", None) is not None else 3,
                    round(float(getattr(sa, "source_tier_multiplier", 1.0)), 4),
                    round(float(getattr(sa, "recency", 0.0)), 4),
                    " | ".join(insights),
                ]
            )

        if (
            self.include_degradation_warning
            and degradation_status
            and getattr(degradation_status, "is_degraded", False)
        ):
            warn_ws = wb.create_sheet("Degradation")
            warn_ws.append(["Field", "Value"])
            warn_ws.append(["is_degraded", True])
            warn_ws.append(["success_rate", f"{degradation_status.success_rate:.3f}"])
            warn_ws.append(
                [
                    "failed/total",
                    f"{degradation_status.failed_attempts}/{degradation_status.total_attempts}",
                ]
            )
            for w in degradation_status.warnings:
                warn_ws.append(["warning", w])

        wb.save(filepath)
        return str(filepath)
